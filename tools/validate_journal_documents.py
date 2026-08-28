"""Validate six current resources, OOXML integrity and fidelity to approved sources.

Structural checks do not replace rendering every page. Author input fields and
explicit template instructions are intentional, not unfinished release prose.
"""
import argparse
import hashlib
import json
import re
from pathlib import Path
from urllib.parse import unquote
from zipfile import ZipFile
from docx import Document
from lxml import etree

ROOT = Path(__file__).resolve().parents[1]
CONTRACT = json.loads((ROOT / 'tools/document-contract.json').read_text(encoding='utf-8'))
NS = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
PROHIBITED = ('REPLACE_WITH', 'Lorem ipsum', '10.xxxx', 'example@example', 'John Doe', 'Jane Doe', 'chiatechresearch@gmail.com', 'CHIATECH-STEM')


def validate(downloads):
    errors, summaries = [], []
    expected = {d['filename'] for d in CONTRACT['documents']}
    actual = {p.name for p in downloads.glob('*.docx') if not p.name.startswith('~$')}
    if actual != expected:
        errors.append(f'Download inventory: missing={sorted(expected-actual)}, unexpected={sorted(actual-expected)}')
    for spec in CONTRACT['documents']:
        path = downloads / spec['filename']
        if not path.is_file() or path.stat().st_size < 50000:
            errors.append(f'{path.name}: missing or unexpectedly small'); continue
        try:
            with ZipFile(path) as package:
                if package.testzip(): errors.append(f'{path.name}: corrupt ZIP member')
                names = set(package.namelist())
                for required in ('[Content_Types].xml', 'word/document.xml', 'word/styles.xml', 'word/settings.xml', 'docProps/core.xml'):
                    if required not in names: errors.append(f'{path.name}: missing {required}')
                trees = {n: etree.fromstring(package.read(n)) for n in names if n.endswith(('.xml', '.rels'))}
                content = [t for n,t in trees.items() if n == 'word/document.xml' or re.match(r'word/(header|footer)\d*\.xml', n)]
                text = '\n'.join(' '.join(t.xpath('//w:t/text()', namespaces=NS)) for t in content)
                for phrase in ['CHIATECH JOURNAL', 'SETEHEM', 'chiatechlibrary@gmail.com', '+234 912 954 8007', *spec['required']]:
                    if phrase.lower() not in text.lower(): errors.append(f'{path.name}: missing required content: {phrase}')
                for phrase in PROHIBITED:
                    if phrase.casefold() in text.casefold(): errors.append(f'{path.name}: prohibited sample or stale content: {phrase}')
                if '\ufffd' in text: errors.append(f'{path.name}: replacement character')
                for n,t in trees.items():
                    if t.xpath('//w:ins | //w:del | //w:commentRangeStart | //w:trackRevisions | //w:vanish', namespaces=NS):
                        errors.append(f'{path.name}: hidden/revision/comment content in {n}')
                    if n.endswith('.rels'):
                        for rel in t:
                            if rel.get('TargetMode') == 'External' and not rel.get('Target','').startswith(('https://','mailto:')):
                                errors.append(f'{path.name}: unsafe/non-HTTPS external relationship')
                if any(n.startswith('word/comments') or 'vbaProject' in n or n.startswith('word/embeddings/') for n in names):
                    errors.append(f'{path.name}: comments, macros or embedded objects')
                images=sum(n.startswith('word/media/') for n in names)
                if images < 2: errors.append(f'{path.name}: missing branding images')
                source=json.loads((ROOT/'tools/document-sources'/f"{spec['id']}.json").read_text())
                if set(source) != names: errors.append(f'{path.name}: OOXML source membership mismatch')
                for member,part in source.items():
                    if member in names and package.read(member) != (ROOT/'tools/document-sources'/part).read_bytes():
                        errors.append(f'{path.name}: source mismatch for {member}')
            doc=Document(path)
            if not any(p.style.name.startswith('Heading') for p in doc.paragraphs): errors.append(f'{path.name}: no semantic headings')
            for index,section in enumerate(doc.sections):
                for label,part in [('header',section.header),('footer',section.footer)]:
                    text=' '.join(part._element.xpath('.//w:t/text()'))
                    if 'CHIATECH JOURNAL' not in text: errors.append(f'{path.name}: unbranded {label} {index+1}')
            summaries.append({'filename':path.name,'bytes':path.stat().st_size,'paragraphs':len(doc.paragraphs),'tables':len(doc.tables),'sections':len(doc.sections),'images':images,'sha256':hashlib.sha256(path.read_bytes()).hexdigest()})
        except Exception as exc:
            errors.append(f'{path.name}: package validation failed ({type(exc).__name__})')
    return errors,summaries


def comparison(downloads):
    refs={}
    for directory in ['about','academic-resources','articles','authors','blog','disciplines','ethics','issues','peer-review','policies','portal','review-engine','search','submit','success','downloads']:
        for p in (ROOT/directory).rglob('*.html'):
            if p.relative_to(ROOT).as_posix().startswith(('articles/2026/','articles/JULY 2026/','articles/_editorial_work/')): continue
            for url in re.findall(r'(?:href|src)=["\']([^"\']+\.docx)(?:[?#][^"\']*)?["\']',p.read_text(encoding='utf-8')):
                name=url.rsplit('/',1)[-1]
                refs.setdefault(unquote(name),[]).append(p.relative_to(ROOT).as_posix())
    inventory=(ROOT/'downloads/README.txt').read_text(encoding='utf-8')
    docs=set(re.findall(r'CHIATECH[^\s`/]+\.docx',inventory))
    expected={d['filename'] for d in CONTRACT['documents']}
    actual={p.name for p in downloads.glob('*.docx') if not p.name.startswith('~$')}
    return {'contractVersion':CONTRACT['version'],'publicCount':6,'physicalDownloads':sorted(actual),'htmlReferences':refs,'validatorExpected':sorted(expected),'builderOutputs':{d['filename']:d['builder'] for d in CONTRACT['documents']},'documentationInventory':sorted(docs),'orphanFiles':sorted(actual-set(refs)),'brokenLinks':sorted(set(refs)-actual),'unexpectedFiles':sorted(actual-expected),'missingFiles':sorted(expected-actual),'documentationMismatch':sorted(docs^expected),'authorMirrors':False,'ignoredWordLocks':[p.name for p in downloads.glob('~$*')],'compatibilityAliases':'HTTP redirects in _redirects; no obsolete DOCX copies'}


def main():
    parser=argparse.ArgumentParser(description=__doc__)
    parser.add_argument('--downloads',type=Path,default=ROOT/'downloads')
    parser.add_argument('--report',type=Path)
    args=parser.parse_args()
    errors,summaries=validate(args.downloads)
    report=comparison(args.downloads)
    for key in ['orphanFiles','brokenLinks','unexpectedFiles','missingFiles','documentationMismatch']:
        if report[key]: errors.append(f'{key}: {report[key]}')
    report.update({'documents':summaries,'errors':errors,'visualQA':'Requires separate page renders and review'})
    if args.report:
        args.report.parent.mkdir(parents=True,exist_ok=True)
        args.report.write_text(json.dumps(report,indent=2)+'\n',encoding='utf-8')
    for row in summaries: print(f"CHECKED {row['filename']}: {row['bytes']:,} bytes; {row['tables']} tables; SHA-256 {row['sha256']}")
    if errors:
        print('DOCUMENT VALIDATION FAILED\n'+'\n'.join(errors)); raise SystemExit(1)
    print('PASS: six current DOCX packages, links, inventory, branding, XML, relationships and builder-source fidelity. No author mirrors required.')


if __name__ == '__main__': main()
